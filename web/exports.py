"""DOCX exporters for individual sessions and consolidated baskets.

The reports inside the system are markdown. We render a curated subset of
markdown (headings, bullets, bold, italic, inline code, code blocks,
horizontal rules) into Word using python-docx. Anything else falls through
as plain text — good enough for the kind of decision-grade prose the agents
produce, and it avoids dragging in pandoc.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Mini markdown → docx
# ---------------------------------------------------------------------------

_INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*)"   # bold
    r"|(\*[^*]+\*)"      # italic
    r"|(`[^`]+`)"        # inline code
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Split ``text`` into runs honoring **bold**, *italic*, and `code`."""
    pos = 0
    for m in _INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _render_markdown(doc, text: str) -> None:
    """Append rendered markdown blocks to ``doc``."""
    if not text:
        return
    lines = text.splitlines()
    in_code = False
    for raw in lines:
        line = raw.rstrip()

        # fenced code block
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Menlo"
            run.font.size = Pt(10)
            continue

        if not line.strip():
            doc.add_paragraph()
            continue

        # horizontal rule
        if re.match(r"^\s*(---+|\*\*\*+|___+)\s*$", line):
            p = doc.add_paragraph()
            p.add_run("─" * 40)
            continue

        # headings
        if line.startswith("#"):
            m = re.match(r"^(#+)\s+(.*)$", line)
            if m:
                level = min(len(m.group(1)), 4)
                doc.add_heading(m.group(2).strip(), level=level)
                continue

        # bullets
        bullet_match = re.match(r"^(\s*)[-*•]\s+(.*)$", line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            indent_level = min(indent_spaces // 2, 3)
            style = "List Bullet" if indent_level == 0 else f"List Bullet {indent_level + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:  # template doesn't define deeper levels
                p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, bullet_match.group(2).strip())
            continue

        # numbered lists
        num_match = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if num_match:
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            _add_inline_runs(p, num_match.group(2).strip())
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, line)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _fmt_ts(ts: Optional[float]) -> str:
    if not ts:
        return "—"
    try:
        return datetime.fromtimestamp(float(ts)).strftime("%Y-%m-%d %H:%M:%S")
    except (TypeError, ValueError):
        return "—"


def _fmt_duration(seconds: Optional[float]) -> str:
    if seconds is None:
        return "—"
    s = float(seconds)
    if s < 1:
        return f"{int(s * 1000)}ms"
    if s < 60:
        return f"{s:.1f}s"
    m = int(s // 60)
    r = int(round(s - m * 60))
    return f"{m}m {r}s"


def _add_kv_table(doc, rows: Iterable[Tuple[str, str]]) -> None:
    rows = list(rows)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = str(k)
        cells[1].text = str(v) if v is not None else "—"
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def _team_timing_rows(team_timings: Dict[str, Dict[str, Any]]) -> List[Tuple[str, str]]:
    if not team_timings:
        return []
    order = ["Analyst Team", "Research Team", "Trading Team", "Risk Management", "Portfolio Management"]
    rows = []
    for team in order + [t for t in team_timings if t not in order]:
        t = team_timings.get(team)
        if not t:
            continue
        rows.append((team, _fmt_duration(t.get("duration_s"))))
    return rows


# ---------------------------------------------------------------------------
# Session export
# ---------------------------------------------------------------------------

# Order in which to render report sections, with display labels.
_SESSION_REPORT_ORDER: List[Tuple[str, str]] = [
    ("final_trade_decision", "Portfolio Manager — Final Decision"),
    ("trader_investment_plan", "Trader — Investment Plan"),
    ("investment_plan", "Research Manager — Investment Plan"),
    ("bull_history", "Bull Researcher"),
    ("bear_history", "Bear Researcher"),
    ("aggressive_history", "Risk Management — Aggressive"),
    ("conservative_history", "Risk Management — Conservative"),
    ("neutral_history", "Risk Management — Neutral"),
    ("market_report", "Market Analyst"),
    ("news_report", "News Analyst"),
    ("sentiment_report", "Social Media Analyst"),
    ("fundamentals_report", "Fundamentals Analyst"),
]


def session_to_docx(session: Dict[str, Any]) -> bytes:
    """Render a single session record as a .docx and return the bytes."""
    doc = Document()
    doc.add_heading(f"{session['ticker']} — {session['analysis_date']}", level=0)

    cfg = session.get("config", {}) or {}
    stats = session.get("stats", {}) or {}
    rows: List[Tuple[str, str]] = [
        ("Status", session.get("status", "—")),
        ("Created", _fmt_ts(session.get("created_at"))),
        ("Started", _fmt_ts(session.get("started_at"))),
        ("Completed", _fmt_ts(session.get("completed_at"))),
        ("Provider", cfg.get("llm_provider", "—")),
        ("Quick model", cfg.get("quick_think_llm", "—")),
        ("Deep model", cfg.get("deep_think_llm", "—")),
        ("Research depth", cfg.get("research_depth", "—")),
        ("Tokens (in / out / Σ)", f"{stats.get('tokens_in', 0):,} / {stats.get('tokens_out', 0):,} / {stats.get('tokens_in', 0) + stats.get('tokens_out', 0):,}"),
        ("LLM calls / tool calls", f"{stats.get('llm_calls', 0)} / {stats.get('tool_calls', 0)}"),
    ]
    if session.get("error"):
        rows.append(("Error", str(session["error"])))
    _add_kv_table(doc, rows)

    timing_rows = _team_timing_rows(session.get("team_timings") or {})
    if timing_rows:
        doc.add_heading("Per-team timings", level=2)
        _add_kv_table(doc, timing_rows)

    sections = session.get("report_sections", {}) or {}
    for key, label in _SESSION_REPORT_ORDER:
        content = (sections.get(key) or "").strip()
        if not content:
            continue
        doc.add_heading(label, level=1)
        _render_markdown(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Batch export
# ---------------------------------------------------------------------------


def batch_to_docx(batch: Dict[str, Any]) -> bytes:
    """Render a batch record (basket) as a .docx and return the bytes."""
    doc = Document()
    doc.add_heading(f"Daily Basket — {batch['analysis_date']}", level=0)

    cfg = batch.get("config", {}) or {}
    totals = batch.get("totals", {}) or {}
    items = batch.get("items", []) or []
    completed = sum(1 for it in items if it.get("status") == "completed")
    rows = [
        ("Status", batch.get("status", "—")),
        ("Created", _fmt_ts(batch.get("created_at"))),
        ("Started", _fmt_ts(batch.get("started_at"))),
        ("Completed", _fmt_ts(batch.get("completed_at"))),
        ("Provider", cfg.get("llm_provider", "—")),
        ("Quick model", cfg.get("quick_think_llm", "—")),
        ("Deep model", cfg.get("deep_think_llm", "—")),
        ("Concurrency", cfg.get("max_concurrency", "—")),
        ("Instruments", f"{completed} completed / {len(items)} total"),
        ("Tokens (in / out / Σ)", f"{totals.get('tokens_in', 0):,} / {totals.get('tokens_out', 0):,} / {totals.get('tokens_in', 0) + totals.get('tokens_out', 0):,}"),
        ("LLM calls / tool calls", f"{totals.get('llm_calls', 0)} / {totals.get('tool_calls', 0)}"),
    ]
    if batch.get("error"):
        rows.append(("Error", str(batch["error"])))
    _add_kv_table(doc, rows)

    # Per-team aggregates across the basket.
    team_totals = batch.get("team_totals") or {}
    if team_totals:
        doc.add_heading("Per-team aggregates", level=2)
        order = ["Analyst Team", "Research Team", "Trading Team", "Risk Management", "Portfolio Management"]
        rows = []
        for team in order + [t for t in team_totals if t not in order]:
            t = team_totals.get(team)
            if not t or not t.get("count"):
                continue
            rows.append((
                team,
                f"Σ {_fmt_duration(t.get('total_s'))}  (avg {_fmt_duration(t.get('avg_s'))} · max {_fmt_duration(t.get('max_s'))} · n={t['count']})",
            ))
        _add_kv_table(doc, rows)

    # Per-instrument summary.
    if items:
        doc.add_heading("Per-instrument summary", level=2)
        table = doc.add_table(rows=1 + len(items), cols=4)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        hdr = table.rows[0].cells
        for i, label in enumerate(["Ticker", "Status", "Tokens", "Decision"]):
            hdr[i].text = label
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for r, it in enumerate(items, start=1):
            cells = table.rows[r].cells
            cells[0].text = str(it.get("ticker", ""))
            cells[1].text = str(it.get("status", ""))
            stats = it.get("stats", {}) or {}
            tk = (stats.get("tokens_in", 0) or 0) + (stats.get("tokens_out", 0) or 0)
            cells[2].text = f"{tk:,}" if tk else "—"
            decision = (it.get("final_decision") or "").strip()
            # Pull just the first non-empty markdown line as a one-shot summary.
            head = next((ln.strip() for ln in decision.splitlines() if ln.strip()), "—")
            cells[3].text = head[:140] + ("…" if len(head) > 140 else "")

    # The big consolidated meta-report.
    report = (batch.get("report") or "").strip()
    if report:
        doc.add_heading("Consolidated report", level=1)
        _render_markdown(doc, report)
    elif batch.get("report_error"):
        doc.add_heading("Consolidated report", level=1)
        doc.add_paragraph(f"Report generation failed: {batch['report_error']}")

    # Per-instrument detail (decision + trader plan).
    if items:
        doc.add_heading("Per-instrument detail", level=1)
        for it in items:
            doc.add_heading(str(it.get("ticker", "")), level=2)
            decision = (it.get("final_decision") or "").strip()
            plan = (it.get("trader_plan") or "").strip()
            if decision:
                doc.add_heading("Final decision", level=3)
                _render_markdown(doc, decision)
            if plan:
                doc.add_heading("Trader plan", level=3)
                _render_markdown(doc, plan)
            if it.get("error"):
                doc.add_heading("Error", level=3)
                doc.add_paragraph(str(it["error"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
